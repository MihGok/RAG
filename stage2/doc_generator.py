"""
stage2/doc_generator.py
────────────────────────
Генерирует DOCX-документ курса из структуры и чанков.
Для каждого шага курса подбирается наиболее релевантный чанк по тегам.
"""

from __future__ import annotations

import logging
import os
from typing import List, Dict, Any, Optional

logger = logging.getLogger(__name__)


def _best_chunk_for_step(
    step_tags: List[str],
    step_title: str,
    chunks: List[Dict[str, Any]],
    used_indices: set,
) -> Optional[Dict[str, Any]]:
    """Выбирает наиболее релевантный чанк по тегам и заголовку."""
    step_tags_lower = {t.lower().strip() for t in step_tags}
    step_words      = set(step_title.lower().split())

    best_idx   = None
    best_score = -1

    for i, chunk in enumerate(chunks):
        chunk_tags  = {t.lower().strip() for t in chunk.get("tags", [])}
        chunk_words = set((chunk.get("final_title", "")).lower().split())

        tag_score   = len(step_tags_lower & chunk_tags)
        title_score = len(step_words & chunk_words) * 0.5
        score       = tag_score + title_score

        # Небольшой штраф за повторное использование
        if i in used_indices:
            score -= 0.3

        if score > best_score:
            best_score = score
            best_idx   = i

    return (chunks[best_idx], best_idx) if best_idx is not None else (None, None)


def generate_course_docx(
    course_structure: Dict[str, Any],
    chunks: List[Dict[str, Any]],
    output_path: str,
    topic: str = "",
) -> str:
    """
    Создаёт DOCX документ курса.

    Args:
        course_structure: результат course_generator (modules, steps, query_texts, tags)
        chunks:           список сгенерированных чанков с текстом
        output_path:      путь для сохранения .docx
        topic:            исходная тема курса

    Returns:
        output_path
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("python-docx не установлен. pip install python-docx")
        return _generate_markdown_fallback(course_structure, chunks, output_path, topic)

    doc = Document()

    # ── Стили ──────────────────────────────────────────────────────────────
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(11)

    # ── Титульная страница ────────────────────────────────────────────────
    course_title = course_structure.get("course_title", f"Курс: {topic}")
    title_para   = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(course_title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True

    if topic:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(f"Тема: {topic}")
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_page_break()

    # ── Содержание (текстовое) ────────────────────────────────────────────
    doc.add_heading("Содержание", level=1)
    modules = course_structure.get("modules", [])

    for m_idx, module in enumerate(modules, 1):
        toc_para = doc.add_paragraph(style="List Bullet")
        toc_run  = toc_para.add_run(f"{m_idx}. {module.get('title', '')}")
        toc_run.font.bold = True

        for s_idx, step in enumerate(module.get("steps", []), 1):
            toc_sub = doc.add_paragraph(style="List Bullet 2")
            toc_sub.add_run(f"  {m_idx}.{s_idx}. {step.get('title', '')}")

    doc.add_page_break()

    # ── Модули ───────────────────────────────────────────────────────────
    used_chunk_indices: set = set()

    for m_idx, module in enumerate(modules, 1):
        doc.add_heading(f"Модуль {m_idx}: {module.get('title', '')}", level=1)

        steps = module.get("steps", [])
        for s_idx, step in enumerate(steps, 1):
            step_title = step.get("title", f"Шаг {s_idx}")
            doc.add_heading(f"{m_idx}.{s_idx}. {step_title}", level=2)

            # Поисковые запросы как подсказка читателю
            query_texts = step.get("query_texts", [])
            if query_texts:
                hint_para = doc.add_paragraph()
                hint_run  = hint_para.add_run("📌 Ключевые вопросы: ")
                hint_run.font.italic = True
                hint_run.font.color.rgb = RGBColor(0x33, 0x66, 0x99)
                hint_para.add_run("; ".join(query_texts[:2]))

            # Подбираем релевантный чанк
            chunk, chunk_idx = _best_chunk_for_step(
                step.get("tags", []),
                step_title,
                chunks,
                used_chunk_indices,
            )

            if chunk:
                used_chunk_indices.add(chunk_idx)

                # Summary
                summary = chunk.get("summary", "")
                if summary:
                    sum_para = doc.add_paragraph()
                    sum_para.add_run("Краткое описание:\n").bold = True
                    sum_para.add_run(summary)

                # Tags
                tags = chunk.get("tags", [])
                if tags:
                    tag_para = doc.add_paragraph()
                    tag_run  = tag_para.add_run("Теги: ")
                    tag_run.font.italic = True
                    tag_para.add_run(", ".join(tags[:10]))

                # Main content
                text = chunk.get("merged_text", "")
                if text:
                    doc.add_paragraph()  # spacing
                    # Split by paragraphs and add
                    for para_text in text[:3500].split("\n\n"):
                        para_text = para_text.strip()
                        if para_text:
                            doc.add_paragraph(para_text)
            else:
                doc.add_paragraph("(материал по данной теме не найден в базе знаний)")

            doc.add_paragraph()  # visual spacing

        if m_idx < len(modules):
            doc.add_page_break()

    # ── Итог ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("Заключение", level=1)
    total_steps = sum(len(m.get("steps", [])) for m in modules)
    doc.add_paragraph(
        f"Данный курс охватывает тему «{topic or course_title}» "
        f"в {len(modules)} модулях, включающих {total_steps} учебных шагов. "
        f"Материал основан на анализе открытых учебных курсов с платформы Stepik."
    )

    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    logger.info("DOCX сохранён: %s", output_path)
    return output_path


def _generate_markdown_fallback(
    course_structure: Dict,
    chunks: List[Dict],
    output_path: str,
    topic: str,
) -> str:
    """Fallback: генерирует Markdown если python-docx недоступен."""
    md_path = output_path.replace(".docx", ".md")
    lines   = []

    title = course_structure.get("course_title", f"Курс: {topic}")
    lines.append(f"# {title}\n")

    chunk_map = {c.get("final_title", ""): c for c in chunks}

    for m_idx, module in enumerate(course_structure.get("modules", []), 1):
        lines.append(f"\n## Модуль {m_idx}: {module.get('title', '')}\n")
        for s_idx, step in enumerate(module.get("steps", []), 1):
            lines.append(f"\n### {m_idx}.{s_idx}. {step.get('title', '')}\n")
            # Find chunk by tags
            step_tags = set(step.get("tags", []))
            best = None
            best_score = -1
            for chunk in chunks:
                score = len(step_tags & set(chunk.get("tags", [])))
                if score > best_score:
                    best_score = score
                    best = chunk
            if best:
                lines.append(f"**Summary:** {best.get('summary','')}\n")
                lines.append(best.get("merged_text", "")[:2000])
            lines.append("\n---\n")

    os.makedirs(os.path.dirname(md_path), exist_ok=True)
    with open(md_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    logger.warning("DOCX недоступен, сохранён Markdown: %s", md_path)
    return md_path